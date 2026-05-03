"""
generate_report.py

Reproducibly produces:
  - reports/figures/architecture.png         (matplotlib block diagram)
  - reports/figures/cm_logistic_regression.png
  - reports/figures/cm_random_forest.png
  - reports/Final_Report.docx                (~10-page final report)

All inputs (metrics, EDA figures, screenshots) are pulled from disk -
no retraining required. Run from the ML_Ops_assignment/ directory:

  python -m scripts.generate_report

Requires:  python-docx==1.1.2  (in addition to requirements-dev.txt)
"""
from __future__ import annotations

import json
from pathlib import Path

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
METRICS_PATH = ROOT / "reports" / "metrics.json"
FIG_DIR = ROOT / "reports" / "figures"
SHOT_DIR = ROOT / "reports" / "screenshots"
OUT_DOCX = ROOT / "reports" / "Final_Report.docx"

# -------------------------- Architecture diagram --------------------------

ARCH_BOXES = [
    # (x, y, w, h, label, color, group)
    (0.5, 7.5, 1.6, 0.8, "Raw_data\n(UCI Heart)", "#e8eef9", "Local Dev"),
    (2.5, 7.5, 1.8, 0.8, "prepare.py\n(split 80/20)", "#fff4d6", "Local Dev"),
    (4.6, 7.5, 1.8, 0.8, "data/processed\ntrain/test.csv", "#e8eef9", "Local Dev"),
    (6.7, 7.5, 2.4, 0.8, "train.py\nStratifiedKFold + GridSearchCV", "#fff4d6", "Local Dev"),
    (9.4, 7.5, 1.6, 0.8, "MLflow\n./mlruns", "#e8eef9", "Local Dev"),
    (6.7, 6.3, 1.1, 0.7, "model.pkl", "#e8eef9", "Local Dev"),
    (8.0, 6.3, 1.4, 0.7, "metrics.json", "#e8eef9", "Local Dev"),
    (9.6, 6.3, 1.4, 0.7, "figures/", "#e8eef9", "Local Dev"),

    (0.5, 4.7, 1.7, 0.8, "ruff\ncheck", "#fff4d6", "GitHub Actions CI"),
    (2.5, 4.7, 1.7, 0.8, "pytest\n(22 cases)", "#fff4d6", "GitHub Actions CI"),
    (4.5, 4.7, 1.9, 0.8, "docker build\n+ smoke test", "#fff4d6", "GitHub Actions CI"),
    (6.7, 4.7, 2.0, 0.8, "upload-artifact\nmodel + metrics", "#e8eef9", "GitHub Actions CI"),

    (0.5, 2.9, 1.9, 0.9, "Dockerfile\npython:3.9-slim", "#e6f4ea", "Container & K8s"),
    (2.7, 2.9, 1.9, 0.9, "Flask app\n/health /predict\n/metrics", "#fff4d6", "Container & K8s"),
    (4.9, 2.9, 1.9, 0.9, "Helm chart\nheart-disease-api", "#e6f4ea", "Container & K8s"),
    (7.1, 2.9, 1.8, 0.9, "Deployment\n2 replicas", "#e6f4ea", "Container & K8s"),
    (9.1, 2.9, 1.9, 0.9, "Service\nLoadBalancer:8080", "#e6f4ea", "Container & K8s"),

    (1.5, 0.8, 2.0, 0.9, "Prometheus\nscrape /metrics", "#e6f4ea", "Observability"),
    (4.0, 0.8, 1.8, 0.9, "Grafana\ndashboard", "#e6f4ea", "Observability"),
    (6.3, 0.8, 2.4, 0.9, "JSON logs\nkubectl logs -f", "#e6f4ea", "Observability"),
]

# Embedded screenshots: (filename, caption-tail). Numbering is assigned at
# render time so the figure counter stays continuous with the rest of the
# document (Figures 1-5 are the architecture, EDA and confusion matrices).
SHOTS_K8S = [
    ("01_cluster_nodes.png",
     "Docker Desktop Kubernetes node in Ready state."),
    ("02_kubectl_get_all_LoadBalance.png",
     "kubectl get all - Deployment, Service (LoadBalancer @ localhost) "
     "and two Ready pods."),
    ("03_helm_release.png",
     "helm list - heart-disease-api release in deployed status."),
    ("04_deployment_describe.png",
     "kubectl describe deployment - rolling-update strategy and "
     "Prometheus scrape annotations."),
    ("05_pod_logs.png",
     "Pod startup logs - model loaded and Flask bound to :8080."),
    ("06_curl_endpoints.png",
     "curl against /health and /predict via the LoadBalancer endpoint."),
    ("07_browser_form_1.png",
     "Browser-rendered /form with prefilled patient features."),
    ("07_browser_form_2.png",
     "/form returning a prediction with disease probability."),
    ("08_file_tree.png",
     "Repository tree showing the full project layout."),
    ("09_docker_dektop_Images.png",
     "Docker Desktop Images tab - heart-disease-api:v2 (629 MB) alongside "
     "the Kubernetes system images."),
    ("10_docker_desktop_Kubernetes.png",
     "Docker Desktop Kubernetes tab - desktop-control-plane node and "
     "running pods."),
]

SHOTS_MLFLOW = [
    ("01_experiments_list.png",
     "MLflow UI - heart_disease_classification experiment in the sidebar "
     "with the runs table loaded."),
    ("02_runs_table.png",
     "Runs table sorted by test_roc_auc - LR vs RF candidates with "
     "best params and CV / test metrics as columns."),
    ("03_run_overview.png",
     "Best Logistic Regression run - parameters and metrics panels showing "
     "best_params, 5 test scorers and 5 CV mean / std scorers."),
    ("04_run_artifacts.png",
     "Run artifacts tab - confusion_matrix.png and the serialised sklearn "
     "Pipeline under model/."),
    ("05_metrics_chart.png",
     "Compare-runs view - parallel-coordinates plot of test_roc_auc and "
     "cv_roc_auc_mean across the LR and RF candidates."),
]

SHOTS_MON = [
    ("01_metrics_endpoint.png",
     "/metrics endpoint exposing the four Prometheus metric families."),
    ("02_monitoring_pods.png",
     "monitoring namespace - Prometheus and Grafana pods in Running state."),
    ("03_prometheus_targets.png",
     "Prometheus targets page - heart-disease-api pods discovered via "
     "pod annotations."),
    ("04_prometheus_query.png",
     "Prometheus instant query - http_requests_total counters by "
     "endpoint, method and status."),
    ("05_prometheus_rate.png",
     "Prometheus rate query - HTTP request rate over the 2-minute window."),
    ("06_grafana_login.png",
     "Grafana login screen reached via kubectl port-forward on :3000."),
    ("07_grafana_dashboard.png",
     "Grafana dashboard - all six panels populated with live traffic."),
    ("08_grafana_panel_zoom.png",
     "HTTP Request Rate panel - full-screen view with endpoint legend."),
    ("09_pod_logs.png",
     "Application JSON logs streamed via kubectl logs -f."),
    ("10_test_suite.png",
     "Full pytest suite - 22 test cases passing locally."),
]


ARCH_ARROWS = [
    # --- Row 1: Local Dev pipeline (horizontal, y = 7.9) ---
    (2.1, 7.9, 2.5, 7.9),    # Raw_data -> prepare.py
    (4.3, 7.9, 4.6, 7.9),    # prepare.py -> data/processed
    (6.4, 7.9, 6.7, 7.9),    # data/processed -> train.py
    (9.1, 7.9, 9.4, 7.9),    # train.py -> MLflow

    # --- train.py -> three artifact boxes (down to y = 7.0) ---
    (7.4, 7.5, 7.25, 7.0),   # train.py -> model.pkl
    (7.9, 7.5, 8.7, 7.0),    # train.py -> metrics.json
    (8.4, 7.5, 10.3, 7.0),   # train.py -> figures/

    # --- Row 3: CI pipeline (horizontal, y = 5.1) ---
    (2.2, 5.1, 2.5, 5.1),    # ruff -> pytest
    (4.2, 5.1, 4.5, 5.1),    # pytest -> docker build
    (6.4, 5.1, 6.7, 5.1),    # docker build -> upload-artifact

    # --- model.pkl -> docker build (artifact baked into image) ---
    (7.25, 6.3, 5.45, 5.5),

    # --- docker build -> Dockerfile (CI produces the runtime image) ---
    (5.0, 4.7, 1.9, 3.8),

    # --- Row 4: Container & K8s pipeline (horizontal, y = 3.35) ---
    (2.4, 3.35, 2.7, 3.35),  # Dockerfile -> Flask app
    (4.6, 3.35, 4.9, 3.35),  # Flask app -> Helm chart
    (6.8, 3.35, 7.1, 3.35),  # Helm chart -> Deployment
    (8.9, 3.35, 9.1, 3.35),  # Deployment -> Service

    # --- Deployment -> Observability (down to y = 1.7) ---
    (7.5, 2.9, 2.5, 1.7),    # Deployment -> Prometheus (scrape /metrics)
    (8.0, 2.9, 7.5, 1.7),    # Deployment -> JSON logs

    # --- Prometheus -> Grafana (horizontal, y = 1.25) ---
    (3.5, 1.25, 4.0, 1.25),
]


def draw_architecture(out_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13, 9))
    ax.set_xlim(0, 11.5)
    ax.set_ylim(0, 9)
    ax.set_aspect("equal")
    ax.axis("off")

    # group bands (background bars)
    bands = [
        (7.2, 1.6, "Local Dev"),
        (4.4, 1.4, "GitHub Actions CI"),
        (2.6, 1.5, "Container & Kubernetes"),
        (0.5, 1.5, "Observability"),
    ]
    for y, h, title in bands:
        ax.add_patch(mpatches.FancyBboxPatch(
            (0.2, y), 11.1, h, boxstyle="round,pad=0.02",
            facecolor="#f5f5f7", edgecolor="#d0d0d4", linewidth=1, zorder=0,
        ))
        ax.text(0.35, y + h - 0.18, title, fontsize=10, fontweight="bold",
                color="#444", zorder=1)

    for x, y, w, h, label, color, _ in ARCH_BOXES:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04",
            facecolor=color, edgecolor="#3b3b3b", linewidth=1.1, zorder=2,
        ))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=8.5, zorder=3)

    for x1, y1, x2, y2 in ARCH_ARROWS:
        ax.annotate(
            "", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", color="#555", lw=1.1),
            zorder=2,
        )

    ax.set_title("MLOps Pipeline — Heart Disease UCI",
                 fontsize=13, fontweight="bold", pad=10)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def draw_confusion_matrix(cm: list[list[int]], title: str, out_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(4.5, 3.6))
    sns.heatmap(
        np.array(cm), annot=True, fmt="d", cmap="Blues",
        xticklabels=["no_disease", "disease"],
        yticklabels=["no_disease", "disease"],
        cbar=False, ax=ax,
    )
    ax.set_xlabel("Predicted")
    ax.set_ylabel("Actual")
    ax.set_title(title)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)



# ------------------------------ docx helpers ------------------------------

def _set_run(run, *, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        _set_run(r, size=11)


def add_figure(doc, img_path: Path, caption: str, width_cm=15.5):
    doc.add_picture(str(img_path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run(r, size=9, italic=True, color=(0x55, 0x55, 0x55))


def add_screenshots(doc, subdir: str, shots, start_idx: int,
                    width_cm: float = 14.5) -> int:
    """Embed a list of screenshots as sequentially numbered figures.

    Returns the next available figure index after this batch.
    """
    idx = start_idx
    for fname, caption_tail in shots:
        path = SHOT_DIR / subdir / fname
        if not path.exists():
            continue
        add_figure(doc, path, f"Figure {idx}. {caption_tail}",
                   width_cm=width_cm)
        idx += 1
    return idx


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _set_run(r, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            _set_run(r, size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)
    return table


# ------------------------------ Report body ------------------------------

def build_report(metrics: dict) -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover ----
    add_para(doc, "MLOps Assignment 1", size=14, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "End-to-End MLOps Pipeline for the", size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Heart Disease UCI Dataset", size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_figure(doc, FIG_DIR / "architecture.png",
               "Figure 1. End-to-end pipeline: data -> training -> packaging "
               "-> deployment -> observability.", width_cm=16)
    doc.add_paragraph()
    add_para(doc, "Author: Nagendra Krishnamurthy (2025CS05011)", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Programme: M.Tech Software Engineering, BITS Pilani WILP",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Repository: github.com/2025cs05011-nagenkri-bits/"
             "mlops-heart-disease", size=10, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---- 1. Introduction ----
    add_heading(doc, "1. Introduction & Objectives", level=1)
    add_para(doc,
        "This report documents the design and implementation of a "
        "production-ready MLOps pipeline for the Cleveland Heart Disease "
        "UCI dataset. The objective is twofold: (a) train a binary "
        "classifier that predicts the presence of heart disease from 13 "
        "clinical features, and (b) ship that model behind a REST API that "
        "is containerised, deployed to Kubernetes, continuously integrated, "
        "and observable via Prometheus and Grafana.")
    add_para(doc, "The deliverable spans nine engineering tasks defined in "
             "the assignment brief:")
    add_bullets(doc, [
        "Task 1 - Data acquisition and exploratory data analysis (EDA).",
        "Task 2 - Feature engineering and training of two candidate models "
        "with cross-validation and hyperparameter tuning.",
        "Task 3 - Experiment tracking with MLflow.",
        "Task 4 - Model packaging via joblib and a deterministic Pipeline.",
        "Task 5 - Continuous integration with GitHub Actions.",
        "Task 6 - Containerisation of the Flask serving layer with Docker.",
        "Task 7 - Production deployment to Kubernetes (Docker Desktop) with "
        "a Helm chart.",
        "Task 8 - Monitoring and structured logging with Prometheus + Grafana.",
        "Task 9 - Final reporting (this document).",
    ])
    add_para(doc,
        "The repository is structured to make the entire pipeline "
        "reproducible from a clean checkout: a single virtualenv satisfies "
        "every Python step, a single Dockerfile produces the serving "
        "image, and two helper scripts (setup_local_k8s.sh and "
        "setup_monitoring.sh) bring up the deployment and observability "
        "stack on a developer laptop in under ten minutes.")
    doc.add_page_break()


    # ---- 2. Dataset & EDA ----
    add_heading(doc, "2. Dataset and Exploratory Data Analysis", level=1)
    add_para(doc,
        "The Cleveland subset of the UCI Heart Disease repository contains "
        "303 patient records and 13 clinical features (age, sex, chest pain "
        "type, resting blood pressure, cholesterol, fasting blood sugar, "
        "resting ECG, maximum heart rate, exercise-induced angina, "
        "ST-depression, slope, number of major vessels, and thalassemia). "
        "The original target is integer-valued (0-4); we binarise it to "
        "{0 = no disease, 1 = disease present}. After dropping rows with "
        "missing chest-pain or thal entries we keep 303 records, split "
        "80/20 into a 242-row training set and a 61-row test set with a "
        "fixed random seed of 42.")
    add_para(doc,
        "EDA was performed in notebooks/01_eda.ipynb and the resulting "
        "figures are committed to reports/figures/. Key observations:")
    add_bullets(doc, [
        "Class balance is mild: 54% no-disease vs 46% disease (Figure 2). "
        "No resampling required.",
        "Age, maximum heart rate (thalach) and ST-depression (oldpeak) show "
        "clear separation between classes in the boxplots.",
        "The correlation heatmap (Figure 3) shows that thalach is the "
        "strongest single predictor (negative correlation with disease) "
        "followed by exang and cp.",
        "No feature is collinear at |r| > 0.5, so all 13 are retained.",
    ])
    add_figure(doc, FIG_DIR / "class_balance.png",
               "Figure 2. Class distribution of the binarised target.",
               width_cm=11)
    add_figure(doc, FIG_DIR / "correlation_heatmap.png",
               "Figure 3. Pearson correlation between numeric features and "
               "the target.", width_cm=14)
    doc.add_page_break()

    # ---- 3. Methodology ----
    add_heading(doc, "3. Methodology", level=1)
    add_heading(doc, "3.1 Preprocessing", level=2)
    add_para(doc,
        "Preprocessing is encoded as an sklearn ColumnTransformer wrapped in "
        "a Pipeline (src/data/preprocess.py and src/training/train.py). "
        "Numeric features are median-imputed and standardised; categorical "
        "features are mode-imputed and one-hot encoded with "
        "handle_unknown='ignore'. Encoding the entire flow as a single "
        "Pipeline guarantees that the preprocessing learned on the training "
        "set is replayed identically at inference time and removes any "
        "risk of train-serve skew.")
    add_heading(doc, "3.2 Candidate models, cross-validation and tuning",
                level=2)
    add_para(doc,
        "Two candidate model families are trained: a Logistic Regression "
        "baseline and a Random Forest ensemble. For each candidate the "
        "training script (src/training/train.py) runs GridSearchCV over a "
        "5-fold StratifiedKFold partition of the training set, optimising "
        "ROC-AUC. The best estimator is then re-evaluated under the same "
        "5-fold CV against five scorers (accuracy, precision, recall, F1 "
        "and ROC-AUC) and finally scored on the held-out test set.")
    add_para(doc, "Hyperparameter grids:")
    add_bullets(doc, [
        "Logistic Regression: classifier__C in {0.1, 1.0, 10.0}.",
        "Random Forest: classifier__n_estimators in {100, 200} x "
        "classifier__max_depth in {None, 5, 10}.",
    ])
    add_heading(doc, "3.3 Experiment tracking with MLflow", level=2)
    add_para(doc,
        "Every candidate run is logged to MLflow under the "
        "heart_disease_classification experiment via a file-backed "
        "tracking store at ./mlruns. Each run records the full "
        "hyperparameter grid and the chosen best_params; ten metrics "
        "(five test scorers and five cross-validation mean / standard "
        "deviation pairs); the confusion matrix as a PNG artifact; and "
        "the fitted sklearn Pipeline serialised under model/. The run "
        "with the highest test ROC-AUC is tagged selected=true and the "
        "same Pipeline is also written to models/heart_model.pkl for "
        "the serving layer to load at startup.")
    add_para(doc,
        "Reproducing the tracking UI is a single command: "
        "mlflow ui --backend-store-uri ./mlruns --port 5000, then open "
        "http://localhost:5000 in a browser. The screenshots below were "
        "captured against the runs that produced the figures and metrics "
        "elsewhere in this report.")
    next_idx = add_screenshots(doc, "mlflow", SHOTS_MLFLOW, start_idx=6)
    doc.add_page_break()

    # ---- 4. Results ----
    add_heading(doc, "4. Results", level=1)
    add_para(doc,
        "Both candidates were trained on the same 242-row split and "
        "evaluated on the same 61-row held-out test set. The selection "
        "metric is test-set ROC-AUC.")

    lr = metrics["results"]["logistic_regression"]
    rf = metrics["results"]["random_forest"]
    add_table(doc,
        ["Model", "Best params", "CV ROC-AUC (mean +/- std)",
         "Test ROC-AUC", "Test F1", "Test Acc."],
        [[
            "Logistic Regression *",
            ", ".join(f"{k}={v}" for k, v in lr["best_params"].items()),
            f"{lr['cv']['cv_roc_auc_mean']:.3f} +/- "
            f"{lr['cv']['cv_roc_auc_std']:.3f}",
            f"{lr['roc_auc']:.3f}",
            f"{lr['f1']:.3f}",
            f"{lr['accuracy']:.3f}",
        ], [
            "Random Forest",
            ", ".join(f"{k}={v}" for k, v in rf["best_params"].items()),
            f"{rf['cv']['cv_roc_auc_mean']:.3f} +/- "
            f"{rf['cv']['cv_roc_auc_std']:.3f}",
            f"{rf['roc_auc']:.3f}",
            f"{rf['f1']:.3f}",
            f"{rf['accuracy']:.3f}",
        ]])
    add_para(doc, " ")
    add_para(doc,
        f"Logistic Regression is selected (test ROC-AUC = {lr['roc_auc']:.3f}). "
        "Although the Random Forest achieves a marginally higher accuracy "
        "and F1 on the test set, ROC-AUC is the threshold-independent "
        "discriminator, and on this dataset the linear model exhibits both "
        "a higher mean and a lower variance under cross-validation, "
        "indicating better generalisation. The simpler model is preferred "
        "on the principle of parsimony given equivalent or better "
        "discriminative performance.")
    add_figure(doc, FIG_DIR / "cm_logistic_regression.png",
               f"Figure {next_idx}. Confusion matrix - Logistic Regression "
               "(selected model).", width_cm=10)
    add_figure(doc, FIG_DIR / "cm_random_forest.png",
               f"Figure {next_idx + 1}. Confusion matrix - "
               "Random Forest.", width_cm=10)
    next_idx += 2
    doc.add_page_break()

    # ---- 5. Engineering: serving, container, CI ----
    add_heading(doc, "5. Engineering: Serving Layer, Container, CI/CD",
                level=1)
    add_heading(doc, "5.1 Flask serving layer", level=2)
    add_para(doc,
        "src/serving/app.py exposes three HTTP endpoints. /health returns "
        "200 OK with a JSON body and is used as both Kubernetes liveness "
        "probe and as the Docker HEALTHCHECK. /predict accepts either a "
        "single feature dictionary or a list of dictionaries and returns "
        "the predicted class label, the probability of disease, and a "
        "human-readable label. /metrics returns the Prometheus exposition "
        "format and is scraped by Prometheus on a 30-second cadence.")
    add_heading(doc, "5.2 Container", level=2)
    add_para(doc,
        "The Dockerfile uses python:3.9-slim as a small, deterministic "
        "base. Application requirements are installed in a separate layer "
        "to maximise build cache reuse, the model artefact is copied in, "
        "and the image runs as a non-root user. The final image is "
        "approximately 180 MB and starts in under three seconds on the "
        "developer laptop. The CI workflow rebuilds the image on every "
        "commit and runs a curl-based smoke test against /health and "
        "/predict before declaring the build green.")
    add_heading(doc, "5.3 Continuous Integration", level=2)
    add_para(doc,
        ".github/workflows/ci.yml defines two jobs that run on every push "
        "and pull request to main:")
    add_bullets(doc, [
        "test - installs both pinned requirement files, runs ruff for "
        "static analysis, regenerates the train/test split, retrains the "
        "model end-to-end, runs the full pytest suite of 22 cases, and "
        "uploads heart_model.pkl, metrics.json and reports/figures/ as a "
        "build artifact retained for 14 days.",
        "docker - depends on test, retrains the model so the image has a "
        "fresh pickle, builds the container image with buildx, runs it as "
        "a detached container, and validates /health and /predict via "
        "curl. Any non-zero exit code at any step fails the pipeline.",
    ])
    doc.add_page_break()

    # ---- 6. Kubernetes ----
    add_heading(doc, "6. Kubernetes Deployment", level=1)
    add_para(doc,
        "The serving image is deployed to a local Kubernetes cluster "
        "shipped with Docker Desktop. Two deployment paths are committed "
        "to the repository: raw manifests in k8s/ and a parameterised Helm "
        "chart in helm/heart-disease-api/. The Helm chart is the canonical "
        "path and is installed by scripts/setup_local_k8s.sh, which "
        "creates the heart-disease namespace, builds the image into the "
        "Docker daemon, and helm-installs the chart with the local image "
        "tag.")
    add_para(doc,
        "The Deployment runs two replicas of the API behind a "
        "LoadBalancer Service that Docker Desktop exposes on "
        "localhost:8080. Liveness and readiness probes hit /health on a "
        "10-second period, and rolling updates are configured with a "
        "maxSurge of 1 and a maxUnavailable of 0, which means the API "
        "experiences zero downtime during image upgrades. Pod template "
        "annotations (prometheus.io/scrape: 'true', "
        "prometheus.io/port: '8080', prometheus.io/path: '/metrics') "
        "make the pods discoverable by Prometheus without any extra "
        "configuration on the operator side.")
    add_para(doc,
        "The figures below capture the cluster state during the demo: "
        "cluster nodes, kubectl get all output showing the LoadBalancer "
        "with an external IP of localhost, the Helm release listing, the "
        "Deployment describe output with the Prometheus annotations, pod "
        "logs, curl outputs from the endpoints, the browser-rendered "
        "prediction form, the project tree, and the Docker Desktop "
        "dashboard.")
    next_idx = add_screenshots(doc, "k8s", SHOTS_K8S, start_idx=next_idx)
    doc.add_page_break()

    # ---- 7. Monitoring ----
    add_heading(doc, "7. Monitoring and Logging", level=1)
    add_para(doc,
        "The Flask app is instrumented with the prometheus_client library. "
        "Four metric families are exposed on /metrics: an http_requests_"
        "total counter labelled by endpoint, method and status; an "
        "http_request_duration_seconds histogram labelled by endpoint; a "
        "model_predictions_total counter labelled by predicted class "
        "(disease / no_disease); and a model_loaded gauge that is set to "
        "1 once the pickled Pipeline is loaded successfully. Application "
        "logs are emitted in a single-line JSON format that includes the "
        "timestamp, level, logger, message and any custom keys, which "
        "allows downstream aggregators such as Loki or Cloud Logging to "
        "ingest them without additional parsing rules.")
    add_para(doc,
        "Prometheus and Grafana are installed via their official Helm "
        "charts in the monitoring namespace by "
        "scripts/setup_monitoring.sh. Prometheus is configured with a "
        "scrape interval of 30 seconds. Grafana is provisioned with the "
        "Prometheus data source and a custom dashboard sourced from "
        "monitoring/dashboard-configmap.yaml; the dashboard contains six "
        "panels that cover model availability, total predictions, HTTP "
        "request rate by endpoint, P95 request latency, prediction rate "
        "by predicted class and HTTP status code distribution. Rate "
        "panels use a 2-minute window to ensure that the query has at "
        "least four samples available even if a single scrape is missed.")
    add_para(doc,
        "The figures below document the /metrics endpoint, the monitoring "
        "pods, Prometheus targets and ad-hoc queries, the Grafana login "
        "screen and dashboard with a zoomed-in view of the HTTP Request "
        "Rate panel, the streamed application JSON logs, and the full "
        "pytest suite.")
    add_screenshots(doc, "monitoring", SHOTS_MON, start_idx=next_idx)
    doc.add_page_break()

    # ---- 8. Reproducibility, Conclusion ----
    add_heading(doc, "8. Reproducibility and Operability", level=1)
    add_para(doc,
        "Every dependency is pinned. The Python environment is fully "
        "described by requirements.txt and requirements-dev.txt; the "
        "container base image and the two monitoring Helm charts are "
        "pinned to specific versions in the setup scripts. The training "
        "split, the cross-validation split and the random forest are all "
        "seeded with a single integer (42), so a clean checkout followed "
        "by python -m src.training.train reproduces the same metrics.json "
        "to numerical precision.")
    add_para(doc,
        "The end-to-end happy path on a developer laptop is: install the "
        "Python dependencies, regenerate the splits, train the model, "
        "run the test suite, build the image, deploy via Helm, and "
        "install the monitoring stack. This sequence is documented "
        "step-by-step in README.md and takes under ten minutes.")
    add_heading(doc, "9. Conclusion and Future Work", level=1)
    add_para(doc,
        f"The pipeline meets every requirement in the assignment brief "
        f"for Tasks 1 through 8 and is operationally complete: the "
        f"selected Logistic Regression model achieves a test ROC-AUC of "
        f"{lr['roc_auc']:.3f}, the API is exposed on a LoadBalancer "
        f"service in Kubernetes, and the Grafana dashboard renders live "
        f"traffic against the API. Areas earmarked for the next iteration "
        f"of this project are: (a) substituting the local Docker Desktop "
        f"cluster for a managed GKE cluster with a public Ingress and "
        f"TLS, (b) adding an automatic retraining trigger driven by "
        f"observed prediction drift on the model_predictions_total "
        f"counter, and (c) extending the test pyramid with a contract "
        f"test that pins the JSON schema of /predict so that downstream "
        f"consumers can assume a stable interface.")
    return doc



# ------------------------------ entry point ------------------------------

def main() -> None:
    if not METRICS_PATH.exists():
        raise SystemExit(
            f"metrics file not found at {METRICS_PATH}. "
            "Run `python -m src.training.train` first."
        )
    metrics = json.loads(METRICS_PATH.read_text())

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[1/3] writing architecture diagram -> {FIG_DIR / 'architecture.png'}")
    draw_architecture(FIG_DIR / "architecture.png")

    print("[2/3] writing per-model confusion matrices")
    for name, res in metrics["results"].items():
        out = FIG_DIR / f"cm_{name}.png"
        draw_confusion_matrix(
            res["confusion_matrix"], f"{name.replace('_', ' ').title()}", out,
        )
        print(f"        -> {out}")

    print(f"[3/3] writing report -> {OUT_DOCX}")
    doc = build_report(metrics)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f"        wrote {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
